"""Day 4 — evidence-pack generation.

Turns the risk-assessed tool list into a polished, audit-style DOCX report:
executive summary, per-tool findings, remediation actions, a simplified
Record of Processing Activities (RoPA) register, and a clause appendix.
"""

from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HIGH_THRESHOLD = 70
MEDIUM_THRESHOLD = 30

REMEDIATION_BY_FLAG = {
    "dormant-access": "Revoke this tool's access — it has not been used in over 180 days. If still needed, require re-authorization with a documented justification.",
    "over-broad-scope": "Reduce the granted permission scope to the minimum required for this tool's function (principle of least privilege), then re-issue access with a narrower OAuth scope.",
    "foreign-hosted": "Confirm a valid cross-border transfer basis exists and register this vendor as a Data Processor under the DPDP Act. Evaluate an India-hosted alternative where one exists.",
    "unknown-hosting": "Determine this tool's actual hosting location before continued use — an unverifiable hosting location is itself a compliance gap.",
    "usage-unknown": "Enable usage/sign-in logging for this tool (e.g. Azure AD Premium sign-in logs) so staleness can actually be verified instead of assumed.",
}


def risk_level(score: int) -> str:
    if score >= HIGH_THRESHOLD:
        return "High"
    if score >= MEDIUM_THRESHOLD:
        return "Medium"
    return "Low"


def _set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def generate_evidence_report(tenant_name: str, tools: list, clauses: list, output_path: str, assessment_mode: str):
    doc = Document()

    title = doc.add_heading("TriNetra — Data Privacy & Compliance Evidence Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Tenant: {tenant_name}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Assessment mode: {assessment_mode}")
    meta_run.italic = True

    disclaimer = doc.add_paragraph()
    d_run = disclaimer.add_run(
        "Generated from representative sample data for hackathon demonstration purposes, not a live tenant scan. "
        "Regulatory citations are an illustrative starter set and require legal verification before real customer use."
    )
    d_run.italic = True
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

    # Executive summary
    doc.add_heading("1. Executive Summary", level=1)
    high = [t for t in tools if risk_level(t["risk_score"]) == "High"]
    medium = [t for t in tools if risk_level(t["risk_score"]) == "Medium"]
    low = [t for t in tools if risk_level(t["risk_score"]) == "Low"]

    doc.add_paragraph(f"Total tools discovered: {len(tools)}")
    summary_table = doc.add_table(rows=2, cols=3)
    summary_table.style = "Light Grid Accent 1"
    for i, label in enumerate(["High Risk", "Medium Risk", "Low Risk"]):
        _set_cell_text(summary_table.rows[0].cells[i], label, bold=True)
    for i, count in enumerate([len(high), len(medium), len(low)]):
        _set_cell_text(summary_table.rows[1].cells[i], str(count))

    doc.add_paragraph()
    doc.add_paragraph(
        f"{len(high)} of {len(tools)} discovered tools require immediate remediation. "
        f"{len(medium)} require scheduled review. {len(low)} show no significant risk indicators at this time."
    )

    # Findings table
    doc.add_heading("2. Findings — All Discovered Tools", level=1)
    findings_table = doc.add_table(rows=1, cols=5)
    findings_table.style = "Light Grid Accent 1"
    headers = ["Tool", "Department", "Risk Level", "Score", "Flags"]
    for i, h in enumerate(headers):
        _set_cell_text(findings_table.rows[0].cells[i], h, bold=True)

    for t in sorted(tools, key=lambda x: -x["risk_score"]):
        row = findings_table.add_row()
        level = risk_level(t["risk_score"])
        color = RGBColor(0xC0, 0x30, 0x30) if level == "High" else (RGBColor(0xB4, 0x53, 0x09) if level == "Medium" else RGBColor(0x2E, 0x7D, 0x32))
        _set_cell_text(row.cells[0], t["tool_name"])
        _set_cell_text(row.cells[1], t["department"])
        _set_cell_text(row.cells[2], level, bold=True, color=color)
        _set_cell_text(row.cells[3], str(t["risk_score"]))
        _set_cell_text(row.cells[4], ", ".join(t["risk_flags"]) if t["risk_flags"] else "—")

    # Detailed findings + remediation for High/Medium
    doc.add_heading("3. Detailed Findings & Remediation", level=1)
    flagged = [t for t in tools if risk_level(t["risk_score"]) in ("High", "Medium")]
    if not flagged:
        doc.add_paragraph("No High or Medium risk tools found.")
    for t in sorted(flagged, key=lambda x: -x["risk_score"]):
        doc.add_heading(f"{t['tool_name']} — {risk_level(t['risk_score'])} ({t['risk_score']}/100)", level=2)
        doc.add_paragraph(t["risk_reasoning"])
        doc.add_paragraph("Recommended action(s):", style="Intense Quote")
        for flag in t["risk_flags"]:
            action = REMEDIATION_BY_FLAG.get(flag, "Review this finding manually.")
            doc.add_paragraph(action, style="List Bullet")

    # RoPA-style register
    doc.add_heading("4. Record of Processing Activities (RoPA) — Simplified Register", level=1)
    doc.add_paragraph(
        "A simplified register mapping each discovered tool to the data it processes, in the spirit of the DPDP Act's "
        "Record of Processing Activities requirement. A production register would include additional fields (legal basis, "
        "retention period, DPO sign-off)."
    )
    ropa_table = doc.add_table(rows=1, cols=4)
    ropa_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Tool", "Data Categories Accessed", "Hosting Region", "Risk Level"]):
        _set_cell_text(ropa_table.rows[0].cells[i], h, bold=True)
    for t in tools:
        row = ropa_table.add_row()
        _set_cell_text(row.cells[0], t["tool_name"])
        _set_cell_text(row.cells[1], ", ".join(t["data_categories_accessed"]))
        _set_cell_text(row.cells[2], t["hosting_region"])
        _set_cell_text(row.cells[3], risk_level(t["risk_score"]))

    # Clause appendix
    doc.add_heading("Appendix A — Regulatory Clauses Referenced", level=1)
    appendix_note = doc.add_paragraph()
    note_run = appendix_note.add_run(
        "Starter/representative clause set — verify against primary legal text before external use."
    )
    note_run.italic = True
    note_run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    for c in clauses:
        p = doc.add_paragraph()
        p.add_run(f"[{c['id']}] {c['framework']} — {c['topic']}: ").bold = True
        p.add_run(c["summary"])

    doc.save(output_path)
    return output_path
